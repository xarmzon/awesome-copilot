#!/usr/bin/env python3
"""
Eyeball - Document analysis with inline source screenshots.

Converts source documents (Word, PDF, web URL) to PDF, renders pages as images,
searches for cited text, highlights matching regions, and assembles an output
Word document with analysis text interleaved with source screenshots.

Usage (called by the Copilot CLI skill, not typically invoked directly):

    python3 eyeball.py build \
        --source <path-or-url> \
        --output <output.docx> \
        --sections '[{"heading": "Section 1", "analysis": "Example analysis text"}]'

    python3 eyeball.py setup-check

    python3 eyeball.py convert --source <file.docx> --output <file.pdf>

    python3 eyeball.py screenshot \
        --source <file.pdf> \
        --anchors '["term1", "term2"]' \
        --page 5 \
        --output screenshot.png
"""

import argparse
import io
import json
import os
import platform
import shutil
import subprocess
import sys
import tempfile

try:
    import fitz  # PyMuPDF
except ImportError:
    fitz = None

try:
    from PIL import Image, ImageDraw
except ImportError:
    Image = None
    ImageDraw = None

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    Document = None
    Inches = None
    Pt = None
    RGBColor = None


def _resolve_path(path_str):
    """Expand ~ and environment variables in a user-provided path."""
    return os.path.expandvars(os.path.expanduser(path_str))


def _check_core_deps():
    """Raise if core dependencies are missing."""
    missing = []
    if fitz is None:
        missing.append("pymupdf")
    if Image is None:
        missing.append("pillow")
    if Document is None:
        missing.append("python-docx")
    if missing:
        print(f"Missing dependencies: {', '.join(missing)}", file=sys.stderr)
        print(f"Run setup.sh or: {sys.executable} -m pip install pymupdf pillow python-docx playwright", file=sys.stderr)
        sys.exit(1)


# ---------------------------------------------------------------------------
# Document conversion: source -> PDF
# ---------------------------------------------------------------------------

def convert_to_pdf(source_path, output_pdf_path):
    """Convert a document to PDF. Supports .docx, .doc, .rtf, .html, .htm."""
    if not os.path.isfile(source_path):
        raise FileNotFoundError(f"Source file not found: {source_path}")

    ext = os.path.splitext(source_path)[1].lower()

    if ext == ".pdf":
        if os.path.abspath(source_path) != os.path.abspath(output_pdf_path):
            shutil.copy2(source_path, output_pdf_path)
        return True

    system = platform.system()

    # Try Microsoft Word first on the current platform
    if system == "Darwin" and ext in (".docx", ".doc", ".rtf"):
        if os.path.exists("/Applications/Microsoft Word.app"):
            if _convert_with_word_mac(source_path, output_pdf_path):
                return True

    if system == "Windows" and ext in (".docx", ".doc", ".rtf"):
        if _convert_with_word_windows(source_path, output_pdf_path):
            return True

    # Fall back to LibreOffice on any platform
    soffice = shutil.which("libreoffice") or shutil.which("soffice")
    if not soffice and system == "Windows":
        soffice = _find_libreoffice_windows()
    if soffice and ext in (".docx", ".doc", ".rtf", ".odt", ".html", ".htm"):
        if _convert_with_libreoffice(soffice, source_path, output_pdf_path):
            return True

    raise RuntimeError(
        f"Cannot convert {ext} to PDF. Install Microsoft Word (macOS/Windows) "
        f"or LibreOffice (any platform)."
    )


def _convert_with_word_mac(source_path, output_pdf_path):
    """Convert using Microsoft Word on macOS via AppleScript."""
    source_abs = os.path.abspath(source_path)
    output_abs = os.path.abspath(output_pdf_path)
    # Escape characters that break AppleScript string interpolation
    source_safe = source_abs.replace('\\', '\\\\').replace('"', '\\"')
    output_safe = output_abs.replace('\\', '\\\\').replace('"', '\\"')
    script = f'''
    tell application "Microsoft Word"
        open POSIX file "{source_safe}"
        delay 5
        set theDoc to active document
        save as theDoc file name POSIX file "{output_safe}" file format format PDF
        close theDoc saving no
    end tell
    '''
    try:
        result = subprocess.run(
            ["osascript", "-e", script],
            capture_output=True, text=True, timeout=120
        )
        return result.returncode == 0 and os.path.exists(output_pdf_path)
    except (subprocess.TimeoutExpired, FileNotFoundError):
        return False


def _convert_with_libreoffice(soffice_path, source_path, output_pdf_path):
    """Convert using LibreOffice headless mode."""
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            result = subprocess.run(
                [soffice_path, "--headless", "--convert-to", "pdf",
                 "--outdir", tmpdir, source_path],
                capture_output=True, text=True, timeout=120
            )
            if result.returncode != 0:
                return False
            basename = os.path.splitext(os.path.basename(source_path))[0] + ".pdf"
            tmp_pdf = os.path.join(tmpdir, basename)
            if os.path.exists(tmp_pdf):
                shutil.move(tmp_pdf, output_pdf_path)
                return True
        except (subprocess.TimeoutExpired, FileNotFoundError):
            pass
    return False


def _find_libreoffice_windows():
    """Find LibreOffice in common Windows install locations."""
    candidates = []
    for env_var in ("ProgramFiles", "ProgramFiles(x86)"):
        base = os.environ.get(env_var)
        if base:
            candidates.append(os.path.join(base, "LibreOffice", "program", "soffice.exe"))
    for path in candidates:
        if os.path.isfile(path):
            return path
    return None


def _convert_with_word_windows(source_path, output_pdf_path):
    """Convert using Microsoft Word on Windows via win32com."""
    word = None
    doc = None
    try:
        import win32com.client
        source_abs = os.path.abspath(source_path)
        output_abs = os.path.abspath(output_pdf_path)
        os.makedirs(os.path.dirname(output_abs), exist_ok=True)

        # DispatchEx creates an isolated Word process; fall back to Dispatch
        # if the DCOM class isn't registered
        try:
            word = win32com.client.DispatchEx("Word.Application")
        except Exception:
            word = win32com.client.Dispatch("Word.Application")

        word.Visible = False
        word.DisplayAlerts = 0
        try:
            word.AutomationSecurity = 3  # msoAutomationSecurityForceDisable
        except Exception:
            pass

        doc = word.Documents.Open(
            FileName=source_abs,
            ConfirmConversions=False,
            ReadOnly=True,
            AddToRecentFiles=False,
            NoEncodingDialog=True,
        )
        doc.ExportAsFixedFormat(
            OutputFileName=output_abs,
            ExportFormat=17,  # wdExportFormatPDF
            OpenAfterExport=False,
        )
        return os.path.isfile(output_abs)
    except Exception:
        return False
    finally:
        if doc is not None:
            try:
                doc.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def render_url_to_pdf(url, output_pdf_path):
    """Render a web page to PDF using Playwright."""
    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        raise RuntimeError(
            "Playwright is required for web URL support. "
            f"Run: {sys.executable} -m pip install playwright && "
            f"{sys.executable} -m playwright install chromium"
        )

    with sync_playwright() as p:
        browser = None
        try:
            browser = p.chromium.launch(headless=True)
            page = browser.new_page()
            page.goto(url, wait_until="networkidle", timeout=30000)

            # Clean up navigation/footer elements for cleaner output
            page.evaluate("""
                document.querySelectorAll(
                    'header, footer, nav, [data-testid="header"], [data-testid="footer"], '
                    + '.site-header, .site-footer, #cookie-banner, .cookie-consent'
                ).forEach(el => el.remove());
            """)

            page.pdf(
                path=output_pdf_path,
                format="Letter",
                print_background=True,
                margin={"top": "0.5in", "bottom": "0.5in",
                        "left": "0.75in", "right": "0.75in"}
            )
        finally:
            if browser is not None:
                browser.close()


# ---------------------------------------------------------------------------
# Screenshot generation
# ---------------------------------------------------------------------------

def screenshot_region(pdf_doc, anchors, target_page=None, target_pages=None,
                      context_padding=40, dpi=200):
    """
    Find anchor text in a PDF and capture the surrounding region as a highlighted image.

    Args:
        pdf_doc: An open fitz.Document.
        anchors: List of search strings. The crop region expands to cover all of them.
        target_page: Single 1-indexed page to search on.
        target_pages: List of 1-indexed pages to search across (results stitched vertically).
        context_padding: Extra padding in PDF points above/below the anchor region.
        dpi: Render resolution.

    Returns:
        (image_bytes, page_label, (width, height)) or (None, None, None).
    """
    if isinstance(anchors, str):
        anchors = [anchors]

    # Determine pages to search
    if target_pages:
        pages = [p - 1 for p in target_pages]
    elif target_page is not None:
        pages = [target_page - 1]
    else:
        pages = list(range(pdf_doc.page_count))

    # Collect hits across pages
    page_hits = {}
    for pg_idx in pages:
        if pg_idx < 0 or pg_idx >= pdf_doc.page_count:
            continue
        page = pdf_doc[pg_idx]
        hits_on_page = []
        for anchor in anchors:
            found = page.search_for(anchor)
            if found:
                hits_on_page.extend([(anchor, h) for h in found])
        if hits_on_page:
            page_hits[pg_idx] = hits_on_page

    if not page_hits:
        return None, None, None

    zoom = dpi / 72

    # If single page, render one region
    if len(page_hits) == 1:
        pg_idx = list(page_hits.keys())[0]
        img = _render_page_region(pdf_doc, pg_idx, page_hits[pg_idx],
                                   context_padding, zoom)
        img_bytes = _img_to_bytes(img)
        return img_bytes, f"page {pg_idx + 1}", img.size

    # Multi-page: stitch vertically
    images = []
    pages_used = sorted(page_hits.keys())
    for pg_idx in pages_used:
        img = _render_page_region(pdf_doc, pg_idx, page_hits[pg_idx],
                                   context_padding, zoom)
        images.append(img)

    stitched = _stitch_vertical(images)
    img_bytes = _img_to_bytes(stitched)

    if len(pages_used) > 1:
        page_nums = ", ".join(str(p + 1) for p in pages_used)
        page_label = f"pages {page_nums}"
    else:
        page_label = f"page {pages_used[0]+1}"

    return img_bytes, page_label, stitched.size


def _render_page_region(pdf_doc, pg_idx, hits_with_anchors, context_padding, zoom):
    """Render a cropped region of a PDF page with highlighted anchor text."""
    page = pdf_doc[pg_idx]
    page_rect = page.rect

    all_rects = [h for _, h in hits_with_anchors]
    min_y = min(r.y0 for r in all_rects)
    max_y = max(r.y1 for r in all_rects)

    crop_rect = fitz.Rect(
        page_rect.x0 + 20,
        max(page_rect.y0, min_y - context_padding),
        page_rect.x1 - 20,
        min(page_rect.y1, max_y + context_padding)
    )

    mat = fitz.Matrix(zoom, zoom)
    pix = page.get_pixmap(matrix=mat, clip=crop_rect)
    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

    # Highlight each anchor hit
    draw = ImageDraw.Draw(img, "RGBA")
    pad = max(2, round(2 * zoom))
    for anchor, rect in hits_with_anchors:
        if rect.y0 >= crop_rect.y0 - 5 and rect.y1 <= crop_rect.y1 + 5:
            x0 = (rect.x0 - crop_rect.x0) * zoom
            y0 = (rect.y0 - crop_rect.y0) * zoom
            x1 = (rect.x1 - crop_rect.x0) * zoom
            y1 = (rect.y1 - crop_rect.y0) * zoom
            draw.rectangle([x0-pad, y0-pad, x1+pad, y1+pad], fill=(255, 255, 0, 100))

    # Border
    ImageDraw.Draw(img).rectangle(
        [0, 0, img.width - 1, img.height - 1],
        outline=(160, 160, 160), width=2
    )

    return img


def _stitch_vertical(images, gap=4):
    """Stitch multiple images vertically with a small gap between them."""
    total_height = sum(img.height for img in images) + gap * (len(images) - 1)
    max_width = max(img.width for img in images)
    stitched = Image.new("RGB", (max_width, total_height), (255, 255, 255))
    y = 0
    for img in images:
        stitched.paste(img, (0, y))
        y += img.height + gap
    ImageDraw.Draw(stitched).rectangle(
        [0, 0, stitched.width - 1, stitched.height - 1],
        outline=(160, 160, 160), width=2
    )
    return stitched


def _img_to_bytes(img):
    """Convert PIL Image to a PNG BytesIO buffer (file-like object)."""
    buf = io.BytesIO()
    img.save(buf, format="PNG")
    buf.seek(0)
    return buf


# ---------------------------------------------------------------------------
# Output document assembly
# ---------------------------------------------------------------------------

def build_analysis_doc(pdf_doc, sections, output_path, title=None, subtitle=None,
                       source_label=None, dpi=200):
    """
    Build a Word document with analysis sections and inline source screenshots.

    Args:
        pdf_doc: An open fitz.Document (the source, already converted to PDF).
        sections: List of dicts, each with:
            - heading (str): Section heading
            - analysis (str): Analysis text
            - anchors (list[str]): Verbatim phrases from source to search and highlight
            - target_page (int, optional): 1-indexed page to search on
            - target_pages (list[int], optional): Multiple pages to search across
            - context_padding (int, optional): Extra padding in PDF points (default 40)
        output_path: Where to save the output .docx file.
        title: Document title.
        subtitle: Document subtitle.
        source_label: Label for the source (e.g., filename or URL).
        dpi: Screenshot resolution.
    """
    doc = Document()

    # Style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title
    if title:
        doc.add_heading(title, level=1)
    if subtitle:
        p = doc.add_paragraph()
        run = p.add_run(subtitle)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(100, 100, 100)
        doc.add_paragraph("")

    # Sections
    for i, section in enumerate(sections):
        heading = section.get("heading", f"Section {i+1}")
        analysis = section.get("analysis", "")
        anchors = section.get("anchors", [])
        target_page = section.get("target_page")
        target_pages = section.get("target_pages")
        padding = section.get("context_padding", 40)

        doc.add_heading(heading, level=2)
        doc.add_paragraph(analysis)

        if anchors:
            img_bytes, page_label, size = screenshot_region(
                pdf_doc, anchors,
                target_page=target_page,
                target_pages=target_pages,
                context_padding=padding,
                dpi=dpi
            )

            if img_bytes:
                # Source label
                p = doc.add_paragraph()
                anchor_text = ", ".join(f'"{a}"' for a in anchors[:3])
                if len(anchors) > 3:
                    anchor_text += f" (+{len(anchors)-3} more)"
                label = f"[Source: {source_label or 'document'}, {page_label}"
                label += f" -- highlighted: {anchor_text}]"
                run = p.add_run(label)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(120, 120, 120)
                run.font.italic = True
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)

                # Screenshot
                doc.add_picture(img_bytes, width=Inches(5.8))
                doc.paragraphs[-1].paragraph_format.space_after = Pt(12)
            else:
                # Anchors not found
                p = doc.add_paragraph()
                run = p.add_run(
                    f"[Screenshot not available: could not find "
                    f"{', '.join(repr(a) for a in anchors)} in the source document]"
                )
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(180, 50, 50)

    # Footer note
    doc.add_paragraph("")
    note = doc.add_paragraph()
    run = note.add_run(
        "Generated by Eyeball. Each screenshot is captured from the source document "
        "with cited text highlighted in yellow. Screenshots are dynamically sized to "
        "cover the full range of text referenced in the analysis. Review the highlighted "
        "source material to verify each assertion."
    )
    run.font.size = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(130, 130, 130)

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# CLI commands
# ---------------------------------------------------------------------------

def cmd_setup_check():
    """Check if all dependencies are available."""
    checks = {
        "PyMuPDF": False,
        "Pillow": False,
        "python-docx": False,
        "Playwright": False,
        "Chromium browser": False,
        "Word (macOS)": False,
        "Word (Windows)": False,
        "LibreOffice": False,
    }

    try:
        import fitz
        checks["PyMuPDF"] = True
    except ImportError:
        pass

    try:
        from PIL import Image
        checks["Pillow"] = True
    except ImportError:
        pass

    try:
        from docx import Document
        checks["python-docx"] = True
    except ImportError:
        pass

    try:
        from playwright.sync_api import sync_playwright
        checks["Playwright"] = True
    except ImportError:
        pass

    # Check Chromium across all platforms
    pw_cache_candidates = []
    system = platform.system()
    if system == "Darwin":
        pw_cache_candidates.append(os.path.expanduser("~/Library/Caches/ms-playwright"))
    if system == "Windows":
        local_app_data = os.environ.get("LOCALAPPDATA", "")
        if local_app_data:
            pw_cache_candidates.append(os.path.join(local_app_data, "ms-playwright"))
    pw_cache_candidates.append(os.path.expanduser("~/.cache/ms-playwright"))
    # Respect PLAYWRIGHT_BROWSERS_PATH
    custom_pw = os.environ.get("PLAYWRIGHT_BROWSERS_PATH")
    if custom_pw and custom_pw != "0":
        pw_cache_candidates.insert(0, custom_pw)
    for pw_cache in pw_cache_candidates:
        if os.path.isdir(pw_cache) and any(
            d.startswith("chromium") for d in os.listdir(pw_cache)
        ):
            checks["Chromium browser"] = True
            break

    # Check converters -- registry/filesystem only, never launch Word
    if system == "Darwin" and os.path.exists("/Applications/Microsoft Word.app"):
        checks["Word (macOS)"] = True

    if system == "Windows":
        try:
            import winreg
            word_reg_paths = [
                (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\WINWORD.EXE"),
                (winreg.HKEY_CURRENT_USER, r"SOFTWARE\Microsoft\Windows\CurrentVersion\App Paths\WINWORD.EXE"),
                (winreg.HKEY_CLASSES_ROOT, r"Word.Application"),
            ]
            for hive, subkey in word_reg_paths:
                try:
                    winreg.OpenKey(hive, subkey)
                    checks["Word (Windows)"] = True
                    break
                except OSError:
                    pass
        except ImportError:
            pass
        # Check if pywin32 is available for Word automation
        if checks["Word (Windows)"]:
            try:
                import win32com.client  # noqa: F401
            except ImportError:
                checks["Word (Windows)"] = False
                print("  Note: Microsoft Word found but pywin32 is not installed.")
                print(f"  Run: {sys.executable} -m pip install pywin32")

    if shutil.which("libreoffice") or shutil.which("soffice"):
        checks["LibreOffice"] = True
    elif system == "Windows":
        if _find_libreoffice_windows():
            checks["LibreOffice"] = True

    print("Eyeball dependency check:")
    all_core = True
    for name, ok in checks.items():
        status = "OK" if ok else "MISSING"
        marker = "+" if ok else "-"
        print(f"  [{marker}] {name}: {status}")
        if name in ("PyMuPDF", "Pillow", "python-docx") and not ok:
            all_core = False

    has_converter = checks["Word (macOS)"] or checks["Word (Windows)"] or checks["LibreOffice"]
    has_web = checks["Playwright"] and checks["Chromium browser"]

    print("")
    print("Source support:")
    print(f"  PDF files:   {'Ready' if all_core else 'Needs: pip3 install pymupdf pillow python-docx'}")
    print(f"  Word docs:   {'Ready' if has_converter else 'Needs: Microsoft Word or LibreOffice'}")
    print(f"  Web URLs:    {'Ready' if has_web else 'Needs: pip3 install playwright && python3 -m playwright install chromium'}")

    return 0 if all_core else 1


def cmd_convert(args):
    """Convert a document to PDF."""
    source = _resolve_path(args.source)
    output = _resolve_path(args.output)

    if source.startswith(("http://", "https://")):
        render_url_to_pdf(source, output)
    else:
        convert_to_pdf(source, output)

    print(f"Converted: {output} ({os.path.getsize(output)} bytes)")


def cmd_screenshot(args):
    """Generate a single screenshot from a PDF."""
    _check_core_deps()
    source = _resolve_path(args.source)

    if not os.path.isfile(source):
        print(f"Source file not found: {source}", file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(source)[1].lower()
    if ext != ".pdf":
        print(f"Source must be a PDF file (got {ext}). "
              f"Use 'convert' to convert other formats first.", file=sys.stderr)
        sys.exit(1)

    anchors = json.loads(args.anchors)
    target_page = args.page
    padding = args.padding
    dpi = args.dpi

    pdf_doc = fitz.open(source)
    try:
        img_bytes, page_label, size = screenshot_region(
            pdf_doc, anchors,
            target_page=target_page,
            context_padding=padding,
            dpi=dpi
        )
    finally:
        pdf_doc.close()

    if img_bytes:
        output = _resolve_path(args.output)
        with open(output, "wb") as f:
            f.write(img_bytes.getvalue())
        print(f"Screenshot saved: {output} ({size[0]}x{size[1]}px, {page_label})")
    else:
        print(f"No matches found for: {anchors}", file=sys.stderr)
        sys.exit(1)


def cmd_build(args):
    """Build a complete analysis document."""
    _check_core_deps()
    source = _resolve_path(args.source)
    output = _resolve_path(args.output)
    sections = json.loads(args.sections)
    title = args.title
    subtitle = args.subtitle
    dpi = args.dpi

    if not source.startswith(("http://", "https://")) and not os.path.isfile(source):
        print(f"Source file not found: {source}", file=sys.stderr)
        sys.exit(1)

    # Determine source type and convert to PDF
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp_pdf = tmp.name

    pdf_doc = None
    try:
        if source.startswith(("http://", "https://")):
            render_url_to_pdf(source, tmp_pdf)
            source_label = source
        elif source.lower().endswith(".pdf"):
            shutil.copy2(source, tmp_pdf)
            source_label = os.path.basename(source)
        else:
            convert_to_pdf(source, tmp_pdf)
            source_label = os.path.basename(source)

        pdf_doc = fitz.open(tmp_pdf)
        build_analysis_doc(
            pdf_doc, sections, output,
            title=title, subtitle=subtitle,
            source_label=source_label,
            dpi=dpi
        )

        size_kb = os.path.getsize(output) / 1024
        print(f"Analysis saved: {output} ({size_kb:.0f} KB)")

    finally:
        if pdf_doc is not None:
            pdf_doc.close()
        if os.path.exists(tmp_pdf):
            os.unlink(tmp_pdf)


def cmd_extract_text(args):
    """Extract text from a source document (for the AI to read before writing analysis)."""
    _check_core_deps()
    source = _resolve_path(args.source)

    if not source.startswith(("http://", "https://")) and not os.path.isfile(source):
        print(f"Source file not found: {source}", file=sys.stderr)
        sys.exit(1)

    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp_pdf = tmp.name

    pdf_doc = None
    try:
        if source.startswith(("http://", "https://")):
            render_url_to_pdf(source, tmp_pdf)
        elif source.lower().endswith(".pdf"):
            shutil.copy2(source, tmp_pdf)
        else:
            convert_to_pdf(source, tmp_pdf)

        pdf_doc = fitz.open(tmp_pdf)
        for i in range(pdf_doc.page_count):
            text = pdf_doc[i].get_text()
            print(f"\n[PAGE {i+1}]")
            print(text)

    finally:
        if pdf_doc is not None:
            pdf_doc.close()
        if os.path.exists(tmp_pdf):
            os.unlink(tmp_pdf)


def main():
    parser = argparse.ArgumentParser(
        description="Eyeball: Document analysis with inline source screenshots"
    )
    sub = parser.add_subparsers(dest="command")

    # setup-check
    sub.add_parser("setup-check", help="Check dependencies")

    # convert
    p_conv = sub.add_parser("convert", help="Convert a document to PDF")
    p_conv.add_argument("--source", required=True)
    p_conv.add_argument("--output", required=True)

    # screenshot
    p_ss = sub.add_parser("screenshot", help="Generate a screenshot from a PDF")
    p_ss.add_argument("--source", required=True, help="PDF file path")
    p_ss.add_argument("--anchors", required=True, help="JSON array of search terms")
    p_ss.add_argument("--page", type=int, help="Target page (1-indexed)")
    p_ss.add_argument("--padding", type=int, default=40)
    p_ss.add_argument("--dpi", type=int, default=200)
    p_ss.add_argument("--output", required=True, help="Output PNG path")

    # build
    p_build = sub.add_parser("build", help="Build analysis document")
    p_build.add_argument("--source", required=True,
                         help="Source document path or URL")
    p_build.add_argument("--output", required=True,
                         help="Output .docx path")
    p_build.add_argument("--sections", required=True,
                         help="JSON array of section objects")
    p_build.add_argument("--title", help="Document title")
    p_build.add_argument("--subtitle", help="Document subtitle")
    p_build.add_argument("--dpi", type=int, default=200)

    # extract-text
    p_text = sub.add_parser("extract-text",
                            help="Extract text from a document (for AI analysis)")
    p_text.add_argument("--source", required=True)

    args = parser.parse_args()

    if args.command == "setup-check":
        sys.exit(cmd_setup_check())
    elif args.command == "convert":
        cmd_convert(args)
    elif args.command == "screenshot":
        cmd_screenshot(args)
    elif args.command == "build":
        cmd_build(args)
    elif args.command == "extract-text":
        cmd_extract_text(args)
    else:
        parser.print_help()
        sys.exit(1)


if __name__ == "__main__":
    main()
